"""Build the canonical LEAI legal .docx files.

These two .docx files are the *ground truth* for LEAI's legal documents.
Edit them in Word for ongoing changes; this script is here only as a one-time
scaffolder and as a recovery option if the docx files are lost or corrupted.

After editing the docx, run regenerate-md-from-docx.sh to refresh the .md
files that the live HTML pages render.

Usage:
    uv run --with python-docx \\
        LEAI/legal/scripts/build_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Tiny formatting helpers
# ---------------------------------------------------------------------------

def add_intro_italic(doc, text):
    """Add an italic line — used for the 'Last updated' caption."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True


def add_para(doc, *segments):
    """Add a paragraph composed of (text, {bold,italic}) segments.

    Each segment is either a plain string (regular text) or a tuple
    ``(text, "bold")`` / ``(text, "italic")`` / ``(text, "bold-italic")``.
    """
    p = doc.add_paragraph()
    for seg in segments:
        if isinstance(seg, str):
            p.add_run(seg)
            continue
        text, fmt = seg
        run = p.add_run(text)
        if "bold" in fmt:
            run.bold = True
        if "italic" in fmt:
            run.italic = True
    return p


def add_bullet(doc, *segments):
    """Add a bullet-list paragraph composed of formatted segments."""
    p = doc.add_paragraph(style="List Bullet")
    for seg in segments:
        if isinstance(seg, str):
            p.add_run(seg)
            continue
        text, fmt = seg
        run = p.add_run(text)
        if "bold" in fmt:
            run.bold = True
        if "italic" in fmt:
            run.italic = True
    return p


def normalize_styles(doc):
    """Make the body text a touch denser so the docs print well."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Privacy Policy
# ---------------------------------------------------------------------------

def build_privacy_policy(out_path: Path) -> None:
    doc = Document()
    normalize_styles(doc)

    doc.add_heading("LEAI Privacy Policy", level=1)
    add_intro_italic(doc, "Last updated: April 2026")

    add_para(
        doc,
        "Your privacy matters. This page explains, in plain language, what LEAI "
        "collects, what it does not collect, who can see it, and what happens to it.",
    )
    add_para(
        doc,
        ("The short version: ", "bold"),
        "LEAI is anonymous by default. If your instructor enables one of the optional "
        "modes (pseudonymous, identified, or Canvas completion credit), the "
        "guarantee changes — and we say so directly in this document and in the chat "
        "interface itself. We do not sell, advertise, or build a profile of you.",
    )

    # ---- Section 1: anonymity modes ---------------------------------------
    doc.add_heading("1. Anonymous by default — but read on", level=2)
    add_para(
        doc,
        "When your instructor creates a survey, they pick an ",
        ("anonymity mode", "italic"),
        ". This is the single most important thing for understanding what LEAI "
        "collects from you. The chat page shows the active mode in a badge near the "
        "top of the screen — look for it before you start typing.",
    )

    add_para(
        doc,
        ("Anonymous mode (default). ", "bold"),
        "No name, email, student ID, institutional login, or profile information is "
        "asked for or stored. Your conversation is tagged with a randomly generated "
        "session code (for example, ",
        ("id_k3m9x2ab", "italic"),
        ") so the system can group your messages within one conversation. The code "
        "is not linked to you as a person. Your instructor cannot tell which "
        "submission came from which student. This is the default mode and the one "
        "we recommend instructors keep.",
    )

    add_para(
        doc,
        ("Pseudonymous / identified mode. ", "bold"),
        "Before the chat begins you are asked to enter an identifier — a real name, "
        "student ID, nickname, or whatever the instructor specifies. Whatever you "
        "type is stored as part of your session record and is visible to your "
        "instructor and any teaching assistants they share access with. Choose "
        "carefully: if you type your real name or your university ID, the "
        "conversation is no longer anonymous to your instructor. If your instructor "
        "allows nicknames, you may use one.",
    )

    add_para(
        doc,
        ("Canvas completion credit. ", "bold"),
        "If your instructor enables completion credit, a short code is shown to you "
        "at the end of your session. You may choose to enter that code into Canvas "
        "(or another LMS) to receive credit for participating. ",
        (
            "If you do this, your Canvas account becomes linked to your specific "
            "session in your instructor's records. ",
            "bold",
        ),
        "That is the trade-off for receiving credit; you can decline to enter the "
        "code if you prefer to remain unlinked.",
    )

    # ---- Section 2: what we store -----------------------------------------
    doc.add_heading("2. What we store", level=2)
    add_para(doc, "For each message you send or receive in LEAI, the backend stores:")
    add_bullet(doc, "The text of the message")
    add_bullet(doc, "Whether the message came from you or from the AI")
    add_bullet(doc, "A timestamp")
    add_bullet(
        doc,
        "The session code, and — in pseudonymous or identified mode — "
        "the identifier you entered",
    )
    add_bullet(doc, "The survey identifier (which course and which prompt)")
    add_bullet(doc, "The name of the AI model used to generate the reply")
    add_bullet(doc, "Your per-session research-consent choice (see Section 6)")
    add_para(
        doc,
        "That is the full record. We do ",
        ("not", "bold"),
        " store device fingerprints, advertising identifiers, or analytics cookies "
        "tied to your identity.",
    )

    # ---- Section 3: IP / logs ---------------------------------------------
    doc.add_heading("3. IP addresses and access logs", level=2)
    add_para(
        doc,
        "LEAI's application code does not log or store your IP address as part of "
        "any feedback record. However, all web traffic to the LEAI backend reaches "
        "it over the public internet, so the underlying cloud platform (currently "
        "Heroku, on Amazon Web Services in the United States) and our application "
        "server may briefly retain network-level access logs that include IP "
        "addresses for operational purposes such as detecting abuse and routing "
        "traffic. These short-lived logs are not joined to feedback content and are "
        "not used for any analysis of feedback.",
    )

    # ---- Section 4: why we collect it -------------------------------------
    doc.add_heading("4. Why we collect it", level=2)
    add_para(
        doc,
        "The whole point of LEAI is to streamline feedback to your instructor ",
        ("while the course is still running", "italic"),
        ". End-of-quarter surveys help future students; they do nothing for you. "
        "LEAI exists to change that: your instructor reads the aggregated feedback "
        "mid-course and uses it to improve your experience in the remaining weeks "
        "of the class. That is the only operational use of your data. We are not "
        "selling it, not advertising to you, and not building a profile of you.",
    )

    # ---- Section 5: who can see it ----------------------------------------
    doc.add_heading("5. Who can see your data", level=2)
    add_bullet(
        doc,
        ("Your course instructor ", "bold"),
        "and any teaching assistants the instructor grants access to (by sharing "
        "the course's access password) can read the feedback submitted to that "
        "course's surveys, in whatever mode the instructor chose. Anyone the "
        "instructor shares the course password with can read everything submitted "
        "to that course; the instructor is responsible for controlling that scope.",
    )
    add_bullet(
        doc,
        ("The GUII Lab research team ", "bold"),
        "maintains the system and has access to the underlying database as needed "
        "to operate the service (fixing bugs, performing backups, and migrating "
        "data).",
    )
    add_bullet(
        doc,
        ("The AI provider ", "bold"),
        "described in Section 7 receives the text of your conversation and the "
        "system prompt your instructor wrote. It does not receive your name, your "
        "session code, or your identifier.",
    )
    add_bullet(
        doc,
        ("No other third parties. ", "bold"),
        "We do not share, sell, or disclose your feedback to advertisers, data "
        "brokers, or any party outside the operational chain described above.",
    )

    # ---- Section 6: optional research use ---------------------------------
    doc.add_heading("6. Optional research use", level=2)
    add_para(
        doc,
        "We are actively improving LEAI so that future students at your institution "
        "(and others) get a better version of it. ",
        (
            "If you check the optional research-consent box on the welcome screen, ",
            "bold",
        ),
        "you give the GUII Lab permission to analyze the text of your messages in "
        "that session for research about how students use AI feedback tools. This "
        "consent is recorded server-side as a per-message flag, so any later "
        "research analysis is restricted to messages where the flag is true. "
        "Published results would only ever report aggregated or de-identified "
        "findings.",
    )
    add_para(
        doc,
        ("If you do not check that box, ", "bold"),
        "your messages are still readable by your instructor for course improvement "
        "(that is the operational purpose of the tool), but they will not be used "
        "in any GUII Lab research analysis.",
    )
    add_para(
        doc,
        "This consent is ",
        ("per session", "bold"),
        ". Each time you start a new feedback session, the box appears again and "
        "you can choose freshly.",
    )

    # ---- Section 7: AI processing -----------------------------------------
    doc.add_heading("7. AI processing", level=2)
    add_para(
        doc,
        "Your messages are sent to a third-party large language model provider "
        "(currently OpenAI) to generate the conversational prompts you see. The "
        "provider processes the text to return a response. We do not send the "
        "provider your name, your institutional ID, your session code, or any "
        "other identifier — only the conversation text and the system prompt your "
        "instructor wrote. Review the provider's own data-handling policies if you "
        "want details on how it processes API traffic. We may switch providers in "
        "the future; if we do, we will update this section.",
    )

    # ---- Section 8: data retention ----------------------------------------
    doc.add_heading("8. Data retention", level=2)
    add_para(
        doc,
        "LEAI is a research prototype and does not yet enforce an automated "
        "retention schedule. Feedback data is retained on the GUII Lab's backend "
        "for the duration of the course and for as long as the lab continues to "
        "operate the service. We are working toward a documented retention and "
        "deletion schedule and will publish it here once it is in place. In the "
        "meantime, instructors and students may request deletion of specific data "
        "as described in Section 9.",
    )

    # ---- Section 9: your rights -------------------------------------------
    doc.add_heading("9. Your rights", level=2)
    add_para(
        doc,
        "In ",
        ("anonymous", "bold"),
        " mode there is no identity attached to your submissions, so there is no "
        "way for us to look up \u201cyour\u201d data and remove it individually. The "
        "most reliable way to keep something out of the dataset is to not submit "
        "it: you can close the tab at any time and any unsent draft is discarded.",
    )
    add_para(
        doc,
        "In ",
        ("pseudonymous, identified, or Canvas-credit", "bold"),
        " modes the situation is different — your submissions can be located. If "
        "you want a specific session removed in those modes, contact your "
        "instructor or write to the GUII Lab. We will honor reasonable removal "
        "requests subject to verification that the request comes from the person "
        "who made the submission.",
    )

    # ---- Section 10: security ---------------------------------------------
    doc.add_heading("10. Security", level=2)
    add_para(
        doc,
        "Data is transmitted over HTTPS and stored in a managed cloud database "
        "operated by the GUII Lab on commercial cloud infrastructure (currently "
        "Heroku Postgres on Amazon Web Services, US region). LEAI is a research "
        "prototype and its security controls are evolving. Course access uses a "
        "shared password which is hashed before being stored. We do not currently "
        "require multi-factor authentication. Treat LEAI as you would an "
        "early-stage academic tool: do not submit anything you would not be willing "
        "for your instructor or the GUII Lab to read.",
    )

    # ---- Section 11: FERPA ------------------------------------------------
    doc.add_heading("11. FERPA and educational records", level=2)
    add_para(
        doc,
        "LEAI is a research and feedback tool operated by the GUII Lab at UC Santa "
        "Cruz. It is not part of UCSC's official student record system. Submissions "
        "made through LEAI are not part of your education record under the Family "
        "Educational Rights and Privacy Act (FERPA), and they are not shared with "
        "the UCSC registrar. If your instructor uses LEAI to award participation "
        "credit through Canvas, only the completion code (and your decision to "
        "enter it) is shared with the LMS — not the contents of your conversation.",
    )

    # ---- Section 12: changes ----------------------------------------------
    doc.add_heading("12. Changes to this policy", level=2)
    add_para(
        doc,
        "We may update this policy as the tool evolves. Material changes — "
        "anything that affects what is collected, who can see it, or how it is "
        "used — will trigger a fresh consent prompt the next time you open a "
        "feedback session, so you can review and re-confirm. The version date at "
        "the top of this page is the version you are agreeing to today.",
    )

    # ---- Section 13: contact ----------------------------------------------
    doc.add_heading("13. Contact", level=2)
    add_para(
        doc,
        "Questions about this privacy policy, or concerns about how LEAI handles "
        "your data, can be directed to the GUII Lab at UC Santa Cruz.",
    )

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# Terms of Use
# ---------------------------------------------------------------------------

def build_terms_of_use(out_path: Path) -> None:
    doc = Document()
    normalize_styles(doc)

    doc.add_heading("LEAI Terms of Use", level=1)
    add_intro_italic(doc, "Last updated: April 2026")

    add_para(
        doc,
        "Welcome to the ",
        ("LEAI Feedback Collector", "bold"),
        ", a tool developed by the Game User Interaction and Intelligence (GUII) "
        "Lab at UC Santa Cruz. By participating in a feedback session, you agree "
        "to the terms below. Please read them before you continue.",
    )

    # ---- Section 1: what this tool is for ---------------------------------
    doc.add_heading("1. What this tool is for", level=2)
    add_para(
        doc,
        "LEAI streamlines feedback collection to your instructor ",
        ("throughout the quarter", "italic"),
        " — not at the very end. The goal is to give your instructor actionable "
        "signals they can act on while the course is still running, so they can "
        "improve your learning experience week by week instead of waiting until "
        "end-of-term evaluations (when it is too late to help you). When you share "
        "feedback here, you are helping your instructor make the course better for "
        "you and your classmates in real time.",
    )

    # ---- Section 2: voluntary participation -------------------------------
    doc.add_heading("2. Your participation is voluntary", level=2)
    add_para(
        doc,
        "Participating in LEAI feedback sessions is ",
        ("entirely voluntary", "bold"),
        ". Choosing not to participate, or stopping at any time, will not affect "
        "your grade or your standing in the course. You may close the tab at any "
        "point, and any in-progress message that you have not sent is discarded. "
        "(If your instructor has enabled Canvas completion credit, see Section 3.)",
    )

    # ---- Section 3: what info you may be asked for ------------------------
    doc.add_heading("3. What information you may be asked for", level=2)
    add_para(
        doc,
        "LEAI's anonymity guarantee depends on the ",
        ("mode", "bold"),
        " your instructor selected. The chat page shows the active mode in a badge "
        "near the top.",
    )
    add_bullet(
        doc,
        ("Anonymous mode (default). ", "bold"),
        "You are not asked for any identifying information. Do not voluntarily "
        "include your name, student ID, classmates' names, or other identifying "
        "details in the body of your messages — the system does not need them, and "
        "they undermine the anonymity design.",
    )
    add_bullet(
        doc,
        ("Pseudonymous / identified mode. ", "bold"),
        "Before the chat begins you are asked to enter an identifier (a real name, "
        "student ID, or nickname, depending on what your instructor specified). "
        "Whatever you type is stored and visible to your instructor.",
    )
    add_bullet(
        doc,
        ("Canvas completion credit. ", "bold"),
        "If your instructor has enabled completion credit, you will be shown a "
        "short code at the end of the session. Entering that code into Canvas "
        "links your Canvas account to your specific session in your instructor's "
        "records. Entering the code is optional.",
    )
    add_para(
        doc,
        "The full data-handling details for each mode are covered in the LEAI "
        "Privacy Policy.",
    )

    # ---- Section 4: appropriate use ---------------------------------------
    doc.add_heading("4. Appropriate use", level=2)
    add_para(
        doc,
        "Please use this tool to share honest, constructive thoughts about the "
        "course experience. Do not use it to:",
    )
    add_bullet(
        doc,
        "Submit content that is harassing, discriminatory, threatening, defamatory, "
        "sexually explicit, or otherwise violates your institution's code of "
        "conduct.",
    )
    add_bullet(
        doc,
        "Share personally identifying information about classmates, instructors, "
        "or third parties without their consent.",
    )
    add_bullet(
        doc,
        "Attempt to disrupt, probe, reverse-engineer, or exploit the system, or to "
        "use it to attack other systems.",
    )
    add_bullet(doc, "Submit content that infringes intellectual property rights.")
    add_para(
        doc,
        "The GUII Lab reserves the right to remove any content that violates these "
        "rules.",
    )

    # ---- Section 5: AI-generated responses --------------------------------
    doc.add_heading("5. AI-generated responses", level=2)
    add_para(
        doc,
        "The replies you receive during the session come from a ",
        ("large language model", "bold"),
        ". They are conversational prompts designed to help you articulate your "
        "feedback — not authoritative statements from your instructor or the "
        "university. LEAI will not:",
    )
    add_bullet(doc, "Grade you")
    add_bullet(doc, "Speak on behalf of your instructor or the university")
    add_bullet(doc, "Provide official academic, legal, medical, or financial advice")
    add_para(
        doc,
        "If you need an official answer, contact your instructor directly through "
        "your course's normal channels.",
    )

    # ---- Section 6: no individual reply -----------------------------------
    doc.add_heading("6. No guarantee of individual reply", level=2)
    add_para(
        doc,
        "In anonymous mode, the instructor cannot tell which submission came from "
        "which student, so there is no way for them to reply to you individually "
        "through this tool. In identified or Canvas-credit modes the instructor ",
        ("can", "italic"),
        " see who you are, but LEAI is still designed for one-way feedback "
        "collection — there is no built-in reply channel. If you need a direct "
        "response from your instructor, use the course's official communication "
        "channels.",
    )

    # ---- Section 7: minimum age -------------------------------------------
    doc.add_heading("7. Minimum age", level=2)
    add_para(
        doc,
        "You must be at least 13 years old to use this tool. LEAI is not intended "
        "for users under 13 and does not knowingly collect data from them. If you "
        "are under 18, your participation is governed by whatever consent "
        "arrangements your course and institution have in place.",
    )

    # ---- Section 8: FERPA -------------------------------------------------
    doc.add_heading("8. FERPA and educational records", level=2)
    add_para(
        doc,
        "LEAI is operated as a research and feedback tool by the GUII Lab. It is "
        "not part of UCSC's official student record system, and submissions made "
        "through LEAI are not part of your education record under the Family "
        "Educational Rights and Privacy Act (FERPA). Even when used for completion "
        "credit through Canvas, only the completion code is shared with the LMS — "
        "not the contents of your conversation.",
    )

    # ---- Section 9: disclaimers / liability -------------------------------
    doc.add_heading("9. Disclaimers and limitation of liability", level=2)
    add_para(
        doc,
        "LEAI is provided ",
        ("as is", "bold"),
        ", without warranty of any kind, as part of the GUII Lab's research and "
        "teaching mission. The GUII Lab and UC Santa Cruz disclaim, to the maximum "
        "extent permitted by law, any implied warranties of merchantability, "
        "fitness for a particular purpose, and non-infringement, and disclaim any "
        "liability for incidental, consequential, or indirect damages arising from "
        "the use of LEAI.",
    )

    # ---- Section 10: changes ----------------------------------------------
    doc.add_heading("10. Changes to these terms", level=2)
    add_para(
        doc,
        "These terms may be updated as the tool evolves. Material changes — "
        "anything that affects what is collected, what is stored, or how it is "
        "used — will trigger a fresh consent prompt the next time you open a "
        "feedback session, so you can review and re-confirm. The version date at "
        "the top of this page is the version you are agreeing to for the current "
        "session.",
    )

    # ---- Section 11: questions --------------------------------------------
    doc.add_heading("11. Questions", level=2)
    add_para(
        doc,
        "For questions about this tool or these terms, contact the GUII Lab at "
        "UC Santa Cruz.",
    )

    doc.save(str(out_path))


def main() -> None:
    legal_dir = Path(__file__).resolve().parent.parent
    privacy_path = legal_dir / "LEAI-Privacy-Policy.docx"
    terms_path = legal_dir / "LEAI-Terms-of-Use.docx"

    build_privacy_policy(privacy_path)
    build_terms_of_use(terms_path)

    print(f"Wrote {privacy_path}")
    print(f"Wrote {terms_path}")


if __name__ == "__main__":
    main()
